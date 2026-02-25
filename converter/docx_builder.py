"""
Word 文档生成模块
将 markdown-it-py 解析的 Token 流转换为格式化的 Word 文档
"""
import io
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from .styles import Fonts, FontSizes, Colors, Spacing, PageLayout


def _rgb_hex(color):
    """将 RGBColor 转换为十六进制字符串（如 'FF0000'）"""
    return f'{color[0]:02X}{color[1]:02X}{color[2]:02X}'


def _clean_text(text):
    """清理文本中 XML 不兼容的控制字符"""
    if not text:
        return text
    # 保留 \t(0x09), \n(0x0A), \r(0x0D)，移除其他控制字符
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)


# 分数占位符正则
FRAC_PATTERN = re.compile(r'⟦FRAC:(.+?):(.+?)⟧')

# Office Math ML 命名空间
MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


def _create_omml_fraction(numerator, denominator):
    """
    创建 OMML (Office Math ML) 分数 XML 元素

    渲染效果为标准数学分数：
       分子
      ────
       分母

    Args:
        numerator: 分子文本
        denominator: 分母文本

    Returns:
        lxml Element 对象，可直接插入段落
    """
    from lxml import etree

    nsmap = {'m': MATH_NS}

    # 构建 oMath 元素
    oMath = etree.SubElement(etree.Element('dummy'), qn('m:oMath'), nsmap=nsmap)

    # 分数元素
    f = etree.SubElement(oMath, qn('m:f'))

    # 分数属性 — bar 类型（带横线）
    fPr = etree.SubElement(f, qn('m:fPr'))
    fType = etree.SubElement(fPr, qn('m:type'))
    fType.set(qn('m:val'), 'bar')

    # 分子
    num = etree.SubElement(f, qn('m:num'))
    num_r = etree.SubElement(num, qn('m:r'))
    num_t = etree.SubElement(num_r, qn('m:t'))
    num_t.text = numerator

    # 分母
    den = etree.SubElement(f, qn('m:den'))
    den_r = etree.SubElement(den, qn('m:r'))
    den_t = etree.SubElement(den_r, qn('m:t'))
    den_t.text = denominator

    return oMath


def _insert_omml_fraction(paragraph, numerator, denominator):
    """将 OMML 分数插入到段落中"""
    oMath = _create_omml_fraction(numerator, denominator)
    paragraph._element.append(oMath)


def _add_text_with_fractions(paragraph, text, set_run_fn=None):
    """
    处理含有分数占位符的文本，将普通文本添加为 run，
    分数占位符替换为 OMML 数学分数。

    Args:
        paragraph: 目标段落
        text: 可能包含分数占位符的文本
        set_run_fn: 可选的 run 样式设置函数 fn(run)

    Returns:
        是否包含分数（True/False）
    """
    if not FRAC_PATTERN.search(text):
        return False

    parts = FRAC_PATTERN.split(text)
    # parts 结构: [前文, 分子1, 分母1, 中间文本, 分子2, 分母2, ...]
    i = 0
    while i < len(parts):
        if i % 3 == 0:
            # 普通文本部分
            if parts[i]:
                run = paragraph.add_run(_clean_text(parts[i]))
                if set_run_fn:
                    set_run_fn(run)
        elif i % 3 == 1:
            # 分子，下一个是分母
            numerator = parts[i]
            denominator = parts[i + 1] if i + 1 < len(parts) else ''
            _insert_omml_fraction(paragraph, numerator, denominator)
            i += 1  # 跳过分母
        i += 1

    return True


class DocxBuilder:
    """将 Markdown Token 流转换为 Word 文档"""

    def __init__(self):
        self.doc = Document()
        self._setup_page()
        self._setup_default_style()
        # 状态变量
        self._list_level = 0
        self._ordered_list = False
        self._list_counter = {}  # 用于跟踪有序列表计数
        self._in_blockquote = False
        self._table_data = []
        self._current_row = []
        self._current_cell_content = []
        self._in_table_header = False
        self._link_url = None

    def _setup_page(self):
        """设置页面布局"""
        section = self.doc.sections[0]
        section.top_margin = PageLayout.TOP_MARGIN
        section.bottom_margin = PageLayout.BOTTOM_MARGIN
        section.left_margin = PageLayout.LEFT_MARGIN
        section.right_margin = PageLayout.RIGHT_MARGIN

    def _setup_default_style(self):
        """设置全局默认样式"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = Fonts.EN_BODY
        font.size = FontSizes.BODY
        font.color.rgb = Colors.BODY
        # 设置中文字体
        style.element.rPr.rFonts.set(qn('w:eastAsia'), Fonts.CN_BODY)
        # 段落格式
        pf = style.paragraph_format
        pf.space_before = Spacing.BODY_BEFORE
        pf.space_after = Spacing.BODY_AFTER
        pf.line_spacing = Spacing.LINE_SPACING

    def _set_run_font(self, run, font_name_en=None, font_name_cn=None,
                      size=None, bold=False, italic=False, color=None, underline=False):
        """统一设置 run 的字体属性"""
        font = run.font
        if font_name_en:
            font.name = font_name_en
        if font_name_cn:
            run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
        if size:
            font.size = size
        font.bold = bold
        font.italic = italic
        if color:
            font.color.rgb = color
        font.underline = underline

    def build(self, tokens: list) -> io.BytesIO:
        """
        将 Token 流转换为 Word 文档

        Args:
            tokens: markdown-it-py 解析出的 Token 列表

        Returns:
            包含 .docx 文件内容的 BytesIO 对象
        """
        i = 0
        while i < len(tokens):
            token = tokens[i]
            i = self._process_token(token, tokens, i)
            i += 1

        # 保存到内存
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _process_token(self, token, tokens, index):
        """处理单个 Token，返回新的索引位置"""
        token_type = token.type

        # ---- 标题 ----
        if token_type == 'heading_open':
            level = int(token.tag[1])  # h1 -> 1
            index = self._handle_heading(tokens, index, level)

        # ---- 段落 ----
        elif token_type == 'paragraph_open':
            index = self._handle_paragraph(tokens, index)

        # ---- 代码块 ----
        elif token_type == 'fence' or token_type == 'code_block':
            self._handle_code_block(token)

        # ---- 无序列表 ----
        elif token_type == 'bullet_list_open':
            self._list_level += 1
            self._ordered_list = False

        elif token_type == 'bullet_list_close':
            self._list_level -= 1
            if self._list_level == 0:
                self._ordered_list = False

        # ---- 有序列表 ----
        elif token_type == 'ordered_list_open':
            self._list_level += 1
            self._ordered_list = True
            self._list_counter[self._list_level] = 0

        elif token_type == 'ordered_list_close':
            if self._list_level in self._list_counter:
                del self._list_counter[self._list_level]
            self._list_level -= 1
            if self._list_level == 0:
                self._ordered_list = False

        # ---- 列表项 ----
        elif token_type == 'list_item_open':
            if self._ordered_list and self._list_level in self._list_counter:
                self._list_counter[self._list_level] += 1

        # ---- 引用块 ----
        elif token_type == 'blockquote_open':
            self._in_blockquote = True

        elif token_type == 'blockquote_close':
            self._in_blockquote = False

        # ---- 表格 ----
        elif token_type == 'table_open':
            self._table_data = []

        elif token_type == 'thead_open':
            self._in_table_header = True

        elif token_type == 'thead_close':
            self._in_table_header = False

        elif token_type == 'tr_open':
            self._current_row = []

        elif token_type == 'tr_close':
            self._table_data.append(self._current_row)

        elif token_type == 'th_open' or token_type == 'td_open':
            self._current_cell_content = []

        elif token_type == 'th_close' or token_type == 'td_close':
            self._current_row.append(self._current_cell_content)

        elif token_type == 'inline' and self._current_row is not None and (
                len(self._table_data) > 0 or self._in_table_header or self._current_row is not None):
            # 如果在表格中，收集单元格内容
            if hasattr(self, '_current_cell_content') and isinstance(self._current_cell_content, list):
                self._current_cell_content.append(token)

        elif token_type == 'table_close':
            self._handle_table()
            self._table_data = []
            self._current_row = []

        # ---- 水平线 ----
        elif token_type == 'hr':
            self._handle_hr()

        return index

    def _handle_heading(self, tokens, index, level):
        """处理标题"""
        # 查找 inline Token（标题内容）
        index += 1
        inline_token = tokens[index]

        heading = self.doc.add_heading(level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 清除默认的 run
        heading.clear()

        # 处理行内内容
        if inline_token.children:
            self._render_inline_runs(heading, inline_token.children, is_heading=True, heading_level=level)
        else:
            run = heading.add_run(_clean_text(inline_token.content))
            self._style_heading_run(run, level)

        # 设置段落间距
        pf = heading.paragraph_format
        pf.space_before = Spacing.HEADING_BEFORE
        pf.space_after = Spacing.HEADING_AFTER

        # 跳过 heading_close
        index += 1
        return index

    def _style_heading_run(self, run, level):
        """为标题 run 设置样式"""
        size = FontSizes.HEADING_MAP.get(level, FontSizes.BODY)
        self._set_run_font(
            run,
            font_name_en=Fonts.EN_HEADING,
            font_name_cn=Fonts.CN_HEADING,
            size=size,
            bold=True,
            color=Colors.HEADING,
        )

    def _handle_paragraph(self, tokens, index):
        """处理段落"""
        # 下一个 Token 应该是 inline
        index += 1
        if index >= len(tokens):
            return index

        inline_token = tokens[index]

        # 检查是否在列表中
        if self._list_level > 0:
            self._handle_list_item(inline_token)
        elif self._in_blockquote:
            self._handle_blockquote(inline_token)
        else:
            # 检查是否在表格中 - 如果有 _current_row 且 _table_data 正在收集
            para = self.doc.add_paragraph()
            if inline_token.children:
                self._render_inline_runs(para, inline_token.children)
            else:
                run = para.add_run(_clean_text(inline_token.content))
                self._set_run_font(run, Fonts.EN_BODY, Fonts.CN_BODY, FontSizes.BODY)

        # 跳过 paragraph_close
        index += 1
        return index

    def _handle_list_item(self, inline_token):
        """处理列表项"""
        indent_level = self._list_level - 1
        prefix = ""

        if self._ordered_list:
            counter = self._list_counter.get(self._list_level, 0)
            prefix = f"{counter}. "
        else:
            bullets = ["•", "◦", "▪", "▸"]
            bullet_char = bullets[min(indent_level, len(bullets) - 1)]
            prefix = f"{bullet_char} "

        para = self.doc.add_paragraph()
        # 设置缩进
        pf = para.paragraph_format
        pf.left_indent = Cm(1.27 * self._list_level)
        pf.first_line_indent = Cm(-0.63)
        pf.space_before = Spacing.LIST_BEFORE
        pf.space_after = Spacing.LIST_AFTER

        # 添加列表符号
        run = para.add_run(prefix)
        self._set_run_font(run, Fonts.EN_BODY, Fonts.CN_BODY, FontSizes.BODY, color=Colors.BODY)

        # 添加内容
        if inline_token.children:
            self._render_inline_runs(para, inline_token.children)
        else:
            run = para.add_run(_clean_text(inline_token.content))
            self._set_run_font(run, Fonts.EN_BODY, Fonts.CN_BODY, FontSizes.BODY)

    def _handle_blockquote(self, inline_token):
        """处理引用块"""
        para = self.doc.add_paragraph()

        # 设置引用样式 - 左侧边框和缩进
        pf = para.paragraph_format
        pf.left_indent = Cm(1.27)
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)

        # 添加左边框（通过XML）
        pPr = para._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:left w:val="single" w:sz="12" w:space="8" w:color="{_rgb_hex(Colors.QUOTE_BORDER)}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        # 添加内容
        if inline_token.children:
            self._render_inline_runs(para, inline_token.children, default_color=Colors.QUOTE_TEXT)
        else:
            run = para.add_run(_clean_text(inline_token.content))
            self._set_run_font(run, Fonts.EN_BODY, Fonts.CN_BODY, FontSizes.BODY,
                               italic=True, color=Colors.QUOTE_TEXT)

    def _handle_code_block(self, token):
        """处理代码块"""
        code = token.content.rstrip('\n')
        language = token.info.strip() if token.info else ""

        # 添加语言标签（如果有）
        if language:
            lang_para = self.doc.add_paragraph()
            lang_run = lang_para.add_run(f"  {language}")
            self._set_run_font(lang_run, Fonts.EN_CODE, Fonts.CN_CODE, FontSizes.SMALL,
                               color=RGBColor(0x99, 0x99, 0x99))
            pf = lang_para.paragraph_format
            pf.space_before = Spacing.CODE_BLOCK_BEFORE
            pf.space_after = Pt(0)

            # 添加底部灰色背景
            self._set_paragraph_shading(lang_para, Colors.CODE_BLOCK_BG)

        lines = code.split('\n')
        for line_idx, line in enumerate(lines):
            para = self.doc.add_paragraph()

            # 设置代码样式
            run = para.add_run(_clean_text(line) if line else " ")  # 空行保留一个空格
            self._set_run_font(run, Fonts.EN_CODE, Fonts.CN_CODE, FontSizes.CODE_BLOCK,
                               color=Colors.CODE_BLOCK_TEXT)

            # 段落格式
            pf = para.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 1.0
            pf.left_indent = Cm(0.5)

            # 设置背景色
            self._set_paragraph_shading(para, Colors.CODE_BLOCK_BG)

        # 代码块后添加间距
        if lines:
            last_para = self.doc.paragraphs[-1]
            last_para.paragraph_format.space_after = Spacing.CODE_BLOCK_AFTER

    def _set_paragraph_shading(self, paragraph, color):
        """为段落设置背景色"""
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{_rgb_hex(color)}" w:val="clear"/>'
        )
        paragraph._element.get_or_add_pPr().append(shading)

    def _handle_table(self):
        """处理表格"""
        if not self._table_data:
            return

        num_rows = len(self._table_data)
        num_cols = max(len(row) for row in self._table_data) if self._table_data else 0

        if num_cols == 0:
            return

        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 设置表格样式
        table.style = self.doc.styles['Table Grid']

        for row_idx, row_data in enumerate(self._table_data):
            row = table.rows[row_idx]
            for col_idx, cell_content_tokens in enumerate(row_data):
                if col_idx >= num_cols:
                    break
                cell = row.cells[col_idx]

                # 清除默认段落
                cell.paragraphs[0].clear()
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # 渲染单元格内容
                for ct in cell_content_tokens:
                    if hasattr(ct, 'children') and ct.children:
                        self._render_inline_runs(para, ct.children)
                    else:
                        content = ct.content if hasattr(ct, 'content') else str(ct)
                        run = para.add_run(_clean_text(content))
                        if row_idx == 0:
                            self._set_run_font(run, Fonts.EN_BODY, Fonts.CN_BODY,
                                               FontSizes.BODY, bold=True,
                                               color=Colors.TABLE_HEADER_TEXT)
                        else:
                            self._set_run_font(run, Fonts.EN_BODY, Fonts.CN_BODY,
                                               FontSizes.BODY, color=Colors.BODY)

                # 表头样式
                if row_idx == 0:
                    self._set_cell_shading(cell, Colors.TABLE_HEADER_BG)
                elif row_idx % 2 == 0:
                    self._set_cell_shading(cell, Colors.TABLE_ALT_BG)

                # 单元格内边距
                self._set_cell_margins(cell, top=60, bottom=60, start=100, end=100)

    def _set_cell_shading(self, cell, color):
        """设置表格单元格背景颜色"""
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{_rgb_hex(color)}" w:val="clear"/>'
        )
        cell._element.get_or_add_tcPr().append(shading)

    def _set_cell_margins(self, cell, top=0, bottom=0, start=0, end=0):
        """设置单元格内边距（单位：twips）"""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'  <w:top w:w="{top}" w:type="dxa"/>'
            f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
            f'  <w:start w:w="{start}" w:type="dxa"/>'
            f'  <w:end w:w="{end}" w:type="dxa"/>'
            f'</w:tcMar>'
        )
        tcPr.append(tcMar)

    def _handle_hr(self):
        """处理水平线"""
        para = self.doc.add_paragraph()
        pPr = para._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(12)

    def _render_inline_runs(self, paragraph, children, is_heading=False,
                            heading_level=None, default_color=None):
        """
        递归渲染行内 Token 子节点到段落中

        Args:
            paragraph: 目标段落
            children: 行内 Token 的子节点列表
            is_heading: 是否是标题
            heading_level: 标题级别
            default_color: 默认文字颜色
        """
        bold = False
        italic = False
        strikethrough = False
        link_url = None

        i = 0
        while i < len(children):
            child = children[i]

            if child.type == 'text':
                content = child.content
                # 检查是否包含分数占位符
                if FRAC_PATTERN.search(content):
                    if is_heading:
                        def style_fn(r, lvl=heading_level, it=italic):
                            self._style_heading_run(r, lvl)
                            if it:
                                r.font.italic = it
                        _add_text_with_fractions(paragraph, content, style_fn)
                    else:
                        color = default_color or Colors.BODY
                        def style_fn(r, b=bold, it=italic, c=color, st=strikethrough):
                            self._set_run_font(
                                r, Fonts.EN_BODY, Fonts.CN_BODY, FontSizes.BODY,
                                bold=b, italic=it, color=c
                            )
                            if st:
                                r.font.strike = True
                        _add_text_with_fractions(paragraph, content, style_fn)
                else:
                    run = paragraph.add_run(_clean_text(content))
                    if is_heading:
                        self._style_heading_run(run, heading_level)
                        if italic:
                            run.font.italic = True
                    else:
                        color = default_color or Colors.BODY
                        self._set_run_font(
                            run, Fonts.EN_BODY, Fonts.CN_BODY, FontSizes.BODY,
                            bold=bold, italic=italic, color=color
                        )
                        if strikethrough:
                            run.font.strike = True
                if link_url:
                    self._add_hyperlink(paragraph, run, link_url)

            elif child.type == 'code_inline':
                run = paragraph.add_run(_clean_text(child.content))
                self._set_run_font(
                    run, Fonts.EN_CODE, Fonts.CN_CODE, FontSizes.CODE,
                    color=Colors.CODE_TEXT
                )
                # 添加浅灰色背景
                rPr = run.element.get_or_add_rPr()
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>'
                )
                rPr.append(shading)

            elif child.type == 'strong_open':
                bold = True
            elif child.type == 'strong_close':
                bold = False

            elif child.type == 'em_open':
                italic = True
            elif child.type == 'em_close':
                italic = False

            elif child.type == 's_open':
                strikethrough = True
            elif child.type == 's_close':
                strikethrough = False

            elif child.type == 'link_open':
                link_url = child.attrs.get('href', '') if child.attrs else ''
            elif child.type == 'link_close':
                link_url = None

            elif child.type == 'softbreak':
                run = paragraph.add_run('\n')
            elif child.type == 'hardbreak':
                run = paragraph.add_run('\n')

            elif child.type == 'image':
                # 图片只显示 alt 文字
                alt = child.attrs.get('alt', '') if child.attrs else ''
                src = child.attrs.get('src', '') if child.attrs else ''
                run = paragraph.add_run(_clean_text(f"[图片: {alt or src}]"))
                self._set_run_font(run, Fonts.EN_BODY, Fonts.CN_BODY, FontSizes.BODY,
                                   italic=True, color=RGBColor(0x99, 0x99, 0x99))

            i += 1

    def _add_hyperlink(self, paragraph, run, url):
        """为 run 添加超链接样式（Word 中的视觉效果）"""
        run.font.color.rgb = Colors.LINK
        run.font.underline = True
        # 在后面添加 URL 注释
        url_run = paragraph.add_run(f" ({url})")
        self._set_run_font(url_run, Fonts.EN_BODY, Fonts.CN_BODY, FontSizes.SMALL,
                           color=RGBColor(0x99, 0x99, 0x99))


def convert_markdown_to_docx(markdown_text: str) -> io.BytesIO:
    """
    将 Markdown 文本转换为 Word 文档

    Args:
        markdown_text: Markdown 格式的文本

    Returns:
        包含 .docx 文件内容的 BytesIO 对象
    """
    from .md_parser import parse_markdown
    from .latex_converter import convert_latex_in_text

    # 预处理：将 LaTeX 数学表达式转换为 Unicode
    markdown_text = convert_latex_in_text(markdown_text)

    tokens = parse_markdown(markdown_text)
    builder = DocxBuilder()
    return builder.build(tokens)
